#!/usr/bin/env python3
"""
generate_fmea.py — DFMEA / PFMEA report generator

Produces an .xlsx workbook and a .docx narrative report from a structured
list of failure chains (FE -> FM -> FC), following the seven-step AIAG & VDA
(2019) approach: Action Priority (AP) is computed from S/O/D per the rule
table in references/sod_scoring_tables.md — RPN is intentionally not
supported.

Usage:
    python generate_fmea.py --type PFMEA --item "Front Bumper Assembly" \\
        --customer "Acme Motors" --template mechanical-assembly \\
        --chains chains.json --output-dir ./output

    python generate_fmea.py --help

chains.json schema (a JSON list of objects):
[
  {
    "process_step": "Injection Molding",   # PFMEA only, optional for DFMEA
    "category_4m": "Machine",              # PFMEA only: Man/Machine/Material/Method/Environment/Measurement
    "fe": "Failure effect text",
    "fm": "Failure mode text",
    "fc": "Failure cause text",
    "s": 8, "o": 4, "d": 5,                # 1-10 integers; omit s/o/d (or set null)
                                            # to fall back on --auto-fill heuristics
    "pc": "Current prevention control text",
    "dc": "Current detection control text",
    "action": "Optional planned action text",
    "action_owner": "Optional owner",
    "action_due": "Optional due date (YYYY-MM-DD)",
    "action_status": "Optional: Suggested / Decided / Implemented / Closed"
  },
  ...
]
"""

import argparse
import json
import os
import sys
from datetime import date


# ---------------------------------------------------------------------------
# Action Priority logic (S/O/D -> AP), per references/sod_scoring_tables.md
# ---------------------------------------------------------------------------

def _ap_band(s, o, d):
    """Core rule table, no validation. Returns 'H'/'M'/'L'. See sod_scoring_tables.md section 4."""
    if s >= 9:
        if o == 1 and d == 1:
            return "M"
        return "H"

    if 6 <= s <= 8:
        if o >= 6:
            return "H"
        if o in (4, 5):
            return "H" if d >= 5 else "M"
        if o in (2, 3):
            return "M" if d >= 5 else "L"
        return "L"  # o == 1

    if 4 <= s <= 5:
        if o >= 6:
            return "H" if d >= 5 else "M"
        if 2 <= o <= 5:
            return "M" if d >= 5 else "L"
        return "L"  # o == 1

    if 2 <= s <= 3:
        if o >= 4 and d >= 5:
            return "M"
        return "L"

    return "L"  # s == 1


def compute_ap(s, o, d):
    """
    Return ('H'|'M'|'L', borderline: bool) for the given S/O/D (1-10 ints).
    `borderline` is True when nudging O or D by +-1 (within 1-10) would flip
    the AP band — i.e. the chain sits right on a rule boundary and deserves
    a human look rather than blind trust in the computed band.
    """
    for v, name in ((s, "S"), (o, "O"), (d, "D")):
        if not isinstance(v, int) or not (1 <= v <= 10):
            raise ValueError(f"{name} must be an integer 1-10, got {v!r}")

    ap = _ap_band(s, o, d)

    borderline = False
    for delta in (-1, 1):
        if 1 <= o + delta <= 10 and _ap_band(s, o + delta, d) != ap:
            borderline = True
        if 1 <= d + delta <= 10 and _ap_band(s, o, d + delta) != ap:
            borderline = True

    return ap, borderline


def flag_special_characteristic(s, ap):
    """Return 'CC', 'SC', or None per SKILL.md section 5."""
    if s >= 9:
        return "CC"
    if s == 8 and ap in ("H", "M"):
        return "SC"
    return None


# ---------------------------------------------------------------------------
# Auto-fill heuristics — only used when a chain omits s/o/d and --auto-fill
# was passed. Deliberately conservative defaults; never invents field data.
# ---------------------------------------------------------------------------

SEVERITY_KEYWORDS = [
    (10, ["safety", "collision", "injury", "fire", "brake fail", "steering fail"]),
    (9, ["regulatory", "non-compliance", "emissions", "legal"]),
    (8, ["loss of function", "fails to operate", "complete failure", "stops working"]),
    (7, ["degraded", "reduced performance", "intermittent major"]),
    (6, ["secondary function", "convenience feature lost"]),
    (5, ["secondary function degraded"]),
    (4, ["very objectionable", "noticeable defect"]),
    (3, ["objectionable", "cosmetic"]),
    (2, ["slightly objectionable"]),
]


def auto_fill_severity(fe_text):
    text = (fe_text or "").lower()
    for score, keywords in SEVERITY_KEYWORDS:
        if any(k in text for k in keywords):
            return score
    return 5  # neutral default — must be reviewed, never treated as final


def auto_fill_occurrence():
    return 5  # "mature technology, similar application" per SKILL.md 4.1


def auto_fill_detection():
    return 6  # "validated pass/fail test method" per SKILL.md 4.1


def apply_auto_fill(chain):
    if chain.get("s") is None:
        chain["s"] = auto_fill_severity(chain.get("fe", ""))
        chain["_s_auto"] = True
    if chain.get("o") is None:
        chain["o"] = auto_fill_occurrence()
        chain["_o_auto"] = True
    if chain.get("d") is None:
        chain["d"] = auto_fill_detection()
        chain["_d_auto"] = True
    return chain


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

REQUIRED_CHAIN_FIELDS = ("fe", "fm", "fc")


def load_chains(path, auto_fill):
    with open(path, "r", encoding="utf-8") as f:
        chains = json.load(f)

    if not isinstance(chains, list) or not chains:
        raise ValueError("chains file must contain a non-empty JSON list")

    for i, chain in enumerate(chains):
        missing = [k for k in REQUIRED_CHAIN_FIELDS if not chain.get(k)]
        if missing:
            raise ValueError(f"chain #{i + 1} is missing required field(s): {missing}")
        for key in ("s", "o", "d"):
            chain.setdefault(key, None)
        if auto_fill:
            apply_auto_fill(chain)
        for key in ("s", "o", "d"):
            if chain[key] is None:
                raise ValueError(
                    f"chain #{i + 1} is missing '{key}' and --auto-fill was not set. "
                    f"Either provide s/o/d for every chain or pass --auto-fill."
                )
        ap, borderline = compute_ap(chain["s"], chain["o"], chain["d"])
        chain["ap"] = ap
        chain["borderline"] = borderline
        chain["special_char"] = flag_special_characteristic(chain["s"], ap)
        chain.setdefault("pc", "")
        chain.setdefault("dc", "")
        chain.setdefault("process_step", "")
        chain.setdefault("category_4m", "")
        chain.setdefault("action", "")
        chain.setdefault("action_owner", "")
        chain.setdefault("action_due", "")
        chain.setdefault("action_status", "")

    return chains


# ---------------------------------------------------------------------------
# XLSX report
# ---------------------------------------------------------------------------

AP_FILL_COLORS = {"H": "FFC7CE", "M": "FFEB9C", "L": "C6EFCE"}
AP_FONT_COLORS = {"H": "9C0006", "M": "9C6500", "L": "006100"}


def build_xlsx(meta, chains, out_path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = f"{meta['fmea_type']} Report"

    bold = Font(bold=True)
    title_font = Font(bold=True, size=14)
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(bold=True, color="FFFFFF")
    thin = Side(style="thin", color="B7B7B7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    row = 1
    ws.cell(row=row, column=1, value=f"{meta['fmea_type']} — {meta['item_name']}").font = title_font
    row += 1
    header_fields = [
        ("Customer", meta.get("customer", "")),
        ("Project No.", meta.get("project_no", "")),
        ("Team", meta.get("team_members", "")),
        ("Template", meta.get("template", "")),
        ("Date", meta.get("report_date", "")),
    ]
    if meta["fmea_type"] == "DFMEA":
        header_fields += [
            ("System Level", meta.get("system_level", "")),
            ("Design Responsibility", meta.get("design_responsibility", "")),
        ]
    else:
        header_fields += [
            ("Process Name", meta.get("process_name", "")),
            ("Manufacturing Site", meta.get("manufacturing_site", "")),
        ]
    for label, value in header_fields:
        ws.cell(row=row, column=1, value=label).font = bold
        ws.cell(row=row, column=2, value=value)
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="Structure").font = bold
    ws.cell(row=row, column=2, value=meta.get("structure_summary", ""))
    row += 1
    ws.cell(row=row, column=1, value="Function").font = bold
    ws.cell(row=row, column=2, value=meta.get("function_summary", ""))
    row += 2

    columns = [
        ("Process Step", "process_step") if meta["fmea_type"] == "PFMEA" else None,
        ("4M/6M Category", "category_4m") if meta["fmea_type"] == "PFMEA" else None,
        ("Failure Effect (FE)", "fe"),
        ("Failure Mode (FM)", "fm"),
        ("Failure Cause (FC)", "fc"),
        ("S", "s"),
        ("O", "o"),
        ("D", "d"),
        ("AP", "ap"),
        ("Special Char.", "special_char"),
        ("Prevention Control (PC)", "pc"),
        ("Detection Control (DC)", "dc"),
        ("Action", "action"),
        ("Owner", "action_owner"),
        ("Due Date", "action_due"),
        ("Status", "action_status"),
    ]
    columns = [c for c in columns if c is not None]

    table_start_row = row
    for col_idx, (label, _key) in enumerate(columns, start=1):
        cell = ws.cell(row=row, column=col_idx, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")
    row += 1

    for chain in chains:
        for col_idx, (_label, key) in enumerate(columns, start=1):
            value = chain.get(key, "")
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.border = border
            cell.alignment = wrap
            if key == "ap" and value in AP_FILL_COLORS:
                cell.fill = PatternFill("solid", fgColor=AP_FILL_COLORS[value])
                cell.font = Font(color=AP_FONT_COLORS[value], bold=True)
        row += 1

    for col_idx, (label, _key) in enumerate(columns, start=1):
        width = 14
        if label in ("Failure Effect (FE)", "Failure Mode (FM)", "Failure Cause (FC)",
                      "Prevention Control (PC)", "Detection Control (DC)", "Action"):
            width = 32
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.freeze_panes = ws.cell(row=table_start_row + 1, column=1)

    row += 2
    counts = {"H": 0, "M": 0, "L": 0}
    for c in chains:
        counts[c["ap"]] += 1
    ws.cell(row=row, column=1, value="Risk Summary").font = bold
    row += 1
    for level, label in (("H", "High priority"), ("M", "Medium priority"), ("L", "Low priority")):
        ws.cell(row=row, column=1, value=label)
        ws.cell(row=row, column=2, value=counts[level])
        row += 1

    mgmt_review = [c for c in chains if c["s"] >= 9 and c["ap"] in ("H", "M")]
    if mgmt_review:
        row += 1
        ws.cell(row=row, column=1, value="Flagged for management review (S 9-10, AP H/M):").font = bold
        row += 1
        for c in mgmt_review:
            ws.cell(row=row, column=1, value=f"- {c['fm']} (S={c['s']}, AP={c['ap']})")
            row += 1

    wb.save(out_path)


# ---------------------------------------------------------------------------
# DOCX report
# ---------------------------------------------------------------------------

def build_docx(meta, chains, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(f"{meta['fmea_type']} — {meta['item_name']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Customer: {meta.get('customer', '')}    "
        f"Project No.: {meta.get('project_no', '-')}    "
        f"Date: {meta.get('report_date', '')}"
    )
    doc.add_paragraph(f"Team: {meta.get('team_members', '-')}")
    doc.add_paragraph(
        "Methodology: AIAG & VDA FMEA Handbook (2019), seven-step approach. "
        "Risk is expressed as Action Priority (H/M/L), not RPN."
    )

    doc.add_heading("Step 1 — Planning and Preparation", level=1)
    p1 = doc.add_paragraph()
    p1.add_run(f"Item under analysis: ").bold = True
    p1.add_run(meta["item_name"])
    if meta["fmea_type"] == "DFMEA":
        doc.add_paragraph(f"System level: {meta.get('system_level', '-')}")
        doc.add_paragraph(f"Design responsibility: {meta.get('design_responsibility', '-')}")
    else:
        doc.add_paragraph(f"Process: {meta.get('process_name', '-')}")
        doc.add_paragraph(f"Manufacturing site: {meta.get('manufacturing_site', '-')}")

    doc.add_heading("Step 2 — Structure Analysis", level=1)
    doc.add_paragraph(meta.get("structure_summary", "(not provided)"))

    doc.add_heading("Step 3 — Function Analysis", level=1)
    doc.add_paragraph(meta.get("function_summary", "(not provided)"))

    doc.add_heading("Step 4 — Failure Analysis", level=1)
    doc.add_paragraph(
        f"{len(chains)} failure chain(s) identified, each expressed as "
        "Failure Effect (FE) -> Failure Mode (FM) -> Failure Cause (FC)."
    )
    for i, c in enumerate(chains, start=1):
        doc.add_heading(f"Chain {i}: {c['fm']}", level=2)
        if c.get("process_step"):
            doc.add_paragraph(f"Process step: {c['process_step']}  |  Category: {c.get('category_4m', '-')}")
        doc.add_paragraph(f"Failure Effect (FE): {c['fe']}")
        doc.add_paragraph(f"Failure Mode (FM): {c['fm']}")
        doc.add_paragraph(f"Failure Cause (FC): {c['fc']}")

    doc.add_heading("Step 5 — Risk Analysis", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for idx, label in enumerate(["Failure Mode", "S", "O", "D", "AP", "PC", "DC"]):
        hdr[idx].text = label
    for c in chains:
        cells = table.add_row().cells
        cells[0].text = c["fm"]
        cells[1].text = str(c["s"])
        cells[2].text = str(c["o"])
        cells[3].text = str(c["d"])
        cells[4].text = c["ap"]
        cells[5].text = c.get("pc", "")
        cells[6].text = c.get("dc", "")
        if c["ap"] == "H":
            for run in cells[4].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
                run.font.bold = True

    mgmt_review = [c for c in chains if c["s"] >= 9 and c["ap"] in ("H", "M")]
    if mgmt_review:
        p = doc.add_paragraph()
        p.add_run("Flagged for management review (S 9-10, AP H/M): ").bold = True
        p.add_run(", ".join(c["fm"] for c in mgmt_review))

    doc.add_heading("Step 6 — Optimization", level=1)
    actionable = [c for c in chains if c["action"] or c["ap"] in ("H", "M")]
    if actionable:
        atable = doc.add_table(rows=1, cols=5)
        atable.style = "Light Grid Accent 1"
        ahdr = atable.rows[0].cells
        for idx, label in enumerate(["Failure Mode", "Action", "Owner", "Due Date", "Status"]):
            ahdr[idx].text = label
        for c in actionable:
            cells = atable.add_row().cells
            cells[0].text = c["fm"]
            cells[1].text = c.get("action") or "(action to be defined — AP indicates review is warranted)"
            cells[2].text = c.get("action_owner", "")
            cells[3].text = c.get("action_due", "")
            cells[4].text = c.get("action_status", "Suggested")
    else:
        doc.add_paragraph("No chains currently require optimization action (all AP = L).")

    doc.add_heading("Step 7 — Results Documentation", level=1)
    counts = {"H": 0, "M": 0, "L": 0}
    for c in chains:
        counts[c["ap"]] += 1
    doc.add_paragraph(
        f"Risk summary: {counts['H']} High, {counts['M']} Medium, {counts['L']} Low priority chain(s)."
    )
    special = [c for c in chains if c["special_char"]]
    if special:
        p = doc.add_paragraph()
        p.add_run("Special characteristics identified: ").bold = True
        p.add_run(", ".join(f"{c['fm']} ({c['special_char']})" for c in special))
        doc.add_paragraph(
            "These should be carried into the control plan (PFMEA) or design record (DFMEA)."
        )
    else:
        doc.add_paragraph("No special characteristics (CC/SC) identified.")

    doc.save(out_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args(argv=None):
    p = argparse.ArgumentParser(description="Generate a DFMEA/PFMEA report (xlsx + docx).")
    p.add_argument("--type", required=True, choices=["DFMEA", "PFMEA"], dest="fmea_type")
    p.add_argument("--item", required=True, dest="item_name", help="Item / product / process name")
    p.add_argument("--customer", default="")
    p.add_argument("--project-no", default="")
    p.add_argument("--team", default="", dest="team_members")
    p.add_argument("--template", default="generic-fmea")
    p.add_argument("--chains", required=True, help="Path to a JSON file with the failure chains")
    p.add_argument("--output-dir", default="./output")
    p.add_argument("--auto-fill", action="store_true",
                    help="Fill missing s/o/d with conservative heuristic defaults instead of erroring")
    p.add_argument("--structure-summary", default="", help="One-line structure analysis summary")
    p.add_argument("--function-summary", default="", help="One-line function analysis summary")
    # DFMEA-only
    p.add_argument("--system-level", default="")
    p.add_argument("--design-responsibility", default="")
    # PFMEA-only
    p.add_argument("--process-name", default="")
    p.add_argument("--manufacturing-site", default="")
    return p.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    os.makedirs(args.output_dir, exist_ok=True)

    try:
        chains = load_chains(args.chains, args.auto_fill)
    except (ValueError, OSError, json.JSONDecodeError) as e:
        print(f"Error loading chains: {e}", file=sys.stderr)
        return 1

    meta = {
        "fmea_type": args.fmea_type,
        "item_name": args.item_name,
        "customer": args.customer,
        "project_no": args.project_no,
        "team_members": args.team_members,
        "template": args.template,
        "report_date": date.today().isoformat(),
        "structure_summary": args.structure_summary,
        "function_summary": args.function_summary,
        "system_level": args.system_level,
        "design_responsibility": args.design_responsibility,
        "process_name": args.process_name,
        "manufacturing_site": args.manufacturing_site,
    }

    safe_name = "".join(ch if ch.isalnum() or ch in " _-" else "_" for ch in args.item_name).strip().replace(" ", "_")
    xlsx_path = os.path.join(args.output_dir, f"{safe_name}_FMEA.xlsx")
    docx_path = os.path.join(args.output_dir, f"{safe_name}_FMEA.docx")

    try:
        build_xlsx(meta, chains, xlsx_path)
        build_docx(meta, chains, docx_path)
    except ImportError as e:
        print(
            f"Missing dependency: {e}. Install with: pip install openpyxl python-docx",
            file=sys.stderr,
        )
        return 1

    counts = {"H": 0, "M": 0, "L": 0}
    for c in chains:
        counts[c["ap"]] += 1

    print(f"Generated: {xlsx_path}")
    print(f"Generated: {docx_path}")
    print(f"Risk summary: {counts['H']} High, {counts['M']} Medium, {counts['L']} Low")
    borderline = [c for c in chains if c.get("borderline")]
    if borderline:
        print(f"{len(borderline)} chain(s) flagged as borderline AP — recommend manual review.")
    mgmt = [c for c in chains if c["s"] >= 9 and c["ap"] in ("H", "M")]
    if mgmt:
        print(f"{len(mgmt)} chain(s) flagged for management review (S 9-10, AP H/M).")

    return 0


if __name__ == "__main__":
    sys.exit(main())
